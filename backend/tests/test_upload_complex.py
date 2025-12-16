import requests
import os
import json
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_complex_docx(filename="test_complex.docx"):
    doc = Document()
    doc.add_heading('Complex Chemistry Questions', 0)
    
    # Selection Question
    doc.add_paragraph('1. 下列关于化学反应的说法正确的是( )')
    doc.add_paragraph('A. 反应A')
    doc.add_paragraph('B. 反应B')
    
    # Big Question with Sub-questions
    doc.add_paragraph('26. (14分) 这是一个综合实验题。')
    doc.add_paragraph('已知物质X具有性质Y。')
    doc.add_paragraph('(1) 写出X的化学式______。')
    doc.add_paragraph('(2) 解释性质Y的原因______。')
    
    doc.save(filename)
    return filename

def test_upload_complex():
    print("Testing complex file upload...")
    
    filename = create_complex_docx()
    
    try:
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            session = requests.Session()
            session.trust_env = False
            response = session.post(f"{BASE_URL}/upload", files=files)
            
        print(f"Status Code: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Parsed {len(data)} items")
            
            ids = [q['id'] for q in data]
            print(f"IDs: {ids}")
            
            # Checks
            if '1' in ids:
                print("[PASS] Found Selection Question 1")
            else:
                print("[FAIL] Missing Selection Question 1")
                
            if '26_1' in ids and '26_2' in ids:
                print("[PASS] Found Sub-questions 26_1 and 26_2")
            else:
                print(f"[FAIL] Missing sub-questions. Found: {ids}")
                
            # Check content of 26_1
            q26_1 = next((q for q in data if q['id'] == '26_1'), None)
            if q26_1 and "【大题题干】" in q26_1['content'] and "综合实验题" in q26_1['content']:
                 print("[PASS] Content format correct for 26_1")
            else:
                 print(f"[FAIL] Content format incorrect for 26_1: {q26_1['content'] if q26_1 else 'Not Found'}")
                 
        else:
            print(f"[FAIL] Upload failed with status {response.status_code}: {response.text}")
            
    except Exception as e:
        print(f"[FAIL] Exception during upload: {e}")
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_upload_complex()
