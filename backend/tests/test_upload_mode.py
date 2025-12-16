import requests
import os
import json
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_complex_docx(filename="test_mode.docx"):
    doc = Document()
    doc.add_heading('Mode Test Questions', 0)
    
    # Big Question with Sub-questions
    doc.add_paragraph('26. (14分) 这是一个综合实验题。')
    doc.add_paragraph('已知物质X具有性质Y。')
    doc.add_paragraph('(1) 写出X的化学式______。')
    doc.add_paragraph('(2) 解释性质Y的原因______。')
    
    doc.save(filename)
    return filename

def test_upload_modes():
    print("Testing upload modes...")
    filename = create_complex_docx()
    
    session = requests.Session()
    session.trust_env = False
    
    try:
        # 1. Test Sub-question Mode (Default)
        print("\n--- Testing Sub-question Mode ---")
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'mode': 'sub_question'}
            response = session.post(f"{BASE_URL}/upload", files=files, data=data)
        
        if response.status_code == 200:
            res_data = response.json()
            ids = [q['id'] for q in res_data]
            print(f"IDs: {ids}")
            if '26_1' in ids and '26_2' in ids:
                print("[PASS] Sub-questions found")
            else:
                print(f"[FAIL] Expected sub-questions, found: {ids}")
        else:
             print(f"[FAIL] Upload failed: {response.text}")

        # 2. Test Whole Mode
        print("\n--- Testing Whole Mode ---")
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'mode': 'whole'}
            response = session.post(f"{BASE_URL}/upload", files=files, data=data)
            
        if response.status_code == 200:
            res_data = response.json()
            ids = [q['id'] for q in res_data]
            print(f"IDs: {ids}")
            if '26' in ids and '26_1' not in ids:
                print("[PASS] Whole question found (no sub-questions)")
                # Check type
                q26 = next(q for q in res_data if q['id'] == '26')
                print(f"Type: {q26.get('type')}")
                if q26.get('type') == 'big_question_whole':
                    print("[PASS] Type is correct")
                else:
                    print(f"[FAIL] Type incorrect: {q26.get('type')}")
            else:
                print(f"[FAIL] Expected whole question 26, found: {ids}")
        else:
             print(f"[FAIL] Upload failed: {response.text}")

    except Exception as e:
        print(f"[FAIL] Exception: {e}")
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_upload_modes()
