import requests
import os
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_test_docx(filename="test_questions.docx"):
    doc = Document()
    doc.add_heading('Chemistry Questions', 0)
    doc.add_paragraph('1. What is H2O?')
    doc.add_paragraph('2. What is NaCl?')
    doc.save(filename)
    return filename

def test_upload():
    print("Testing file upload...")
    
    # Create a dummy file
    filename = create_test_docx()
    
    try:
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(f"{BASE_URL}/upload", files=files)
            
        print(f"Status Code: {response.status_code}")
        print(f"Response: {response.text}")
        
        if response.status_code == 200:
            print("[PASS] Upload successful")
            data = response.json()
            if len(data) >= 2:
                print(f"[PASS] Parsed {len(data)} questions")
            else:
                print(f"[FAIL] Expected at least 2 questions, got {len(data)}")
        else:
            print(f"[FAIL] Upload failed with status {response.status_code}")
            
    except Exception as e:
        print(f"[FAIL] Exception during upload: {e}")
    finally:
        # Cleanup
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_upload()
