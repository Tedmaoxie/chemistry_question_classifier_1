import requests
import os
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_test_docx(filename="reproduce_test.docx"):
    doc = Document()
    doc.add_heading('Chemistry Questions', 0)
    doc.add_paragraph('1. 下列说法正确的是( )')
    doc.add_paragraph('A. 选项A')
    doc.add_paragraph('B. 选项B')
    doc.save(filename)
    return filename

def test_upload():
    print("Testing file upload...")
    filename = create_test_docx("reproduce_test.DOCX")
    
    try:
        session = requests.Session()
        session.trust_env = False
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'mode': 'sub_question'}
            response = session.post(f"{BASE_URL}/upload", files=files, data=data)
            
        print(f"Status Code: {response.status_code}")
        print(f"Response: {response.text}")
        
    except Exception as e:
        print(f"Exception: {e}")
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_upload()
